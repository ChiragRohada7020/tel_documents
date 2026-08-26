"""
DOCX processor - extracts text from DOCX files.
"""

import logging
from typing import List

import docx

logger = logging.getLogger(__name__)


class DocxProcessor:
    """Service for extracting text content from DOCX files."""

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted text as a single string.
        """
        text = ""
        try:
            doc = docx.Document(file_path)
            paragraphs: List[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Tables frequently contain the important values in forms/invoices.
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            text = "\n".join(paragraphs)
            logger.info(f"Extracted {len(text)} characters from DOCX: {file_path}")
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        return text

    def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """
        Extract text from tables in a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            A list of tables, each table is a list of rows, each row is a list of cell texts.
        """
        tables = []
        try:
            doc = docx.Document(file_path)
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            logger.info(f"Extracted {len(tables)} tables from DOCX: {file_path}")
        except Exception as e:
            logger.error(f"Error extracting tables from DOCX {file_path}: {e}")
        return tables
